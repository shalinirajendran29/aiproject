import os
import shutil
import hashlib
from datetime import datetime
from fastapi import APIRouter, Depends, UploadFile, File, BackgroundTasks, HTTPException
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from typing import Dict, Any, List, Optional
from ..database import get_db
from ..config import settings
from ..models.document import DBModelDocument
from ..models.mapping import DBModelMappingMemory
from ..schemas.document import DocumentResponse, DocumentReview
from ..services.preprocessor import ImagePreprocessor
from ..services.ocr_engine import OCREngine
from ..services.slm_engine import SLMEngine
from ..services.cache_service import cache_service

router = APIRouter(prefix="/documents", tags=["documents"])

# Shared services instances
ocr_engine = OCREngine()
slm_engine = SLMEngine()

def finalize_document_processing(doc_id: str, status: str, extracted_json: Optional[Dict[str, Any]], confidence_score: Optional[float], db: Session):
    """Updates document status, caches the result, updates job stage, and triggers webhook."""
    doc = db.query(DBModelDocument).filter(DBModelDocument.id == doc_id).first()
    if not doc:
        return
        
    doc.status = status
    if extracted_json is not None:
        doc.extracted_json = extracted_json
    if confidence_score is not None:
        doc.confidence_score = float(confidence_score)
    db.commit()

    # Cache document result for 10 minutes (Point 14)
    doc_data = {
        "id": doc.id,
        "filename": doc.filename,
        "storage_path": doc.storage_path,
        "mime_type": doc.mime_type,
        "status": doc.status,
        "extracted_json": doc.extracted_json,
        "corrected_json": doc.corrected_json,
        "confidence_score": float(doc.confidence_score) if doc.confidence_score is not None else None,
        "ocr_raw_text": doc.ocr_raw_text,
        "created_at": doc.created_at.isoformat() if doc.created_at else None
    }
    cache_service.set(f"doc:result:{doc.id}", doc_data, expire_seconds=600)
    
    # Set job progress stage
    cache_service.set(f"job:{doc.id}:stage", "completed" if status == "completed" else "failed", expire_seconds=3600)

    # Webhook Reliability (Point 8)
    from .admin import load_admin_data
    try:
        admin_data = load_admin_data()
        webhook_url = admin_data.get("settings", {}).get("webhook_url")
        webhook_secret = admin_data.get("settings", {}).get("webhook_secret")
        if webhook_url:
            from ..services.webhook_service import dispatch_webhook
            payload = {
                "document_id": doc.id,
                "filename": doc.filename,
                "status": status,
                "extracted_json": doc.extracted_json,
                "confidence_score": float(doc.confidence_score) if doc.confidence_score is not None else 0.0,
                "timestamp": datetime.utcnow().isoformat() + "Z"
            }
            dispatch_webhook(webhook_url, payload, webhook_secret, doc.id)
    except Exception as wh_err:
        print(f"Failed to dispatch webhook for doc {doc.id}: {wh_err}")

def process_document_task(doc_id: str, db_session_maker):
    """Background task to run OpenCV + OCR + SLM pipeline."""
    db = db_session_maker()
    try:
        # 1. Fetch document from db
        doc = db.query(DBModelDocument).filter(DBModelDocument.id == doc_id).first()
        if not doc:
            return
            
        # Simulate Virus scanning quarantine phase
        from .admin import append_admin_log, load_admin_data
        admin_settings = load_admin_data()["settings"]
        
        cache_service.set(f"job:{doc_id}:stage", "quarantine", expire_seconds=3600)
        
        if admin_settings.get("virus_scanning_enabled", True):
            append_admin_log("security", "INFO", f"ClamAV: Starting quarantine virus scan on '{doc.filename}'...")
            # Mock virus scan delay
            import time
            time.sleep(0.4) 
            append_admin_log("security", "INFO", f"ClamAV: Virus scan passed for '{doc.filename}' (clean scan signature).")
        else:
            append_admin_log("security", "WARNING", f"Virus scanner disabled in security settings. Bypassing scan for: {doc.filename}")
            
        doc.status = "processing"
        db.commit()
        
        raw_path = doc.storage_path
        ext = os.path.splitext(raw_path.lower())[1]
        
        # Calculate file hash for OCR and AI caching (Rule 4, 5)
        try:
            with open(raw_path, "rb") as f:
                file_bytes = f.read()
            file_hash = hashlib.sha256(file_bytes).hexdigest()
        except Exception as hash_err:
            print(f"Failed to calculate file hash: {hash_err}")
            file_hash = None
        
        # 1a. Handle Spreadsheet Ingestion
        is_excel = ext.startswith(".xls") or ext in [".ods", ".xslv", ".xlsv"]
        is_csv = ext == ".csv"
        
        if is_excel or is_csv:
            cache_service.set(f"job:{doc_id}:stage", "extraction", expire_seconds=3600)
            import pandas as pd
            
            def clean_excel_key(raw_key: str) -> str:
                import re
                clean = re.sub(r'[^a-zA-Z0-9\s_]', '', str(raw_key)).strip().lower()
                if any(kw in clean for kw in ["code", "id", "identifier", "no", "num", "number"]):
                    if "customer" in clean or "client" in clean or "patient" in clean:
                        return "customer_code"
                    if "pin" in clean or "zip" in clean or "postal" in clean:
                        return "pin_code"
                    if "pan" in clean:
                        return "pan_no"
                    if "gst" in clean or "gstin" in clean:
                        return "gstin_no"
                    if "phone" in clean or "mobile" in clean or "contact" in clean or "tel" in clean:
                        return "mobile_number"
                if "gst" in clean or "gstin" in clean:
                    return "gstin_no"
                if "pan" in clean:
                    return "pan_no"
                if "pin" in clean or "pincode" in clean or "zip" in clean or "postal" in clean:
                    return "pin_code"
                if "state" in clean or "province" in clean or "region" in clean:
                    return "state"
                if "district" in clean or "city" in clean or "town" in clean or "locality" in clean:
                    return "district"
                if "country" in clean or "nation" in clean:
                    return "country"
                if "address" in clean or "addr" in clean or "location" in clean or "residence" in clean or "street" in clean:
                    return "address"
                if "phone" in clean or "mobile" in clean or "contact" in clean or "tel" in clean or "cell" in clean:
                    return "mobile_number"
                if "email" in clean or "mail" in clean:
                    return "email"
                if "dob" in clean or "birth" in clean:
                    return "dob"
                if "name" in clean or "client" in clean or "patient" in clean or "customer" in clean or "employee" in clean or "owner" in clean:
                    return "full_name"
                return re.sub(r'\s+', '_', clean)

            def post_process_record(rec: dict) -> dict:
                processed = {}
                for k, v in rec.items():
                    processed[k] = v
                if "mobile_number" in processed and processed["mobile_number"]:
                    val = str(processed["mobile_number"]).strip()
                    if val.startswith("+91"):
                        val = val[3:].strip()
                    elif val.startswith("91") and len(val) > 10:
                        val = val[2:].strip()
                    val = "".join(c for c in val if c.isdigit())
                    processed["mobile_number"] = val
                if "email" in processed and processed["email"] and "@" in processed["email"]:
                    parts = processed["email"].split("@")
                    if len(parts) == 2:
                        domain = parts[1].strip().replace(" ", ".")
                        domain = domain.rstrip(".,")
                        username = parts[0].strip()
                        processed["email"] = f"{username}@{domain}"
                return processed

            if is_csv:
                df = pd.read_csv(raw_path)
            else:
                if ext not in [".xlsx", ".xls", ".xlsm"]:
                    # Create a temp copy with .xlsx extension to force pandas openpyxl engine
                    temp_xlsx_path = raw_path + "_temp.xlsx"
                    shutil.copyfile(raw_path, temp_xlsx_path)
                    try:
                        df = pd.read_excel(temp_xlsx_path, engine="openpyxl")
                    finally:
                        if os.path.exists(temp_xlsx_path):
                            os.remove(temp_xlsx_path)
                else:
                    df = pd.read_excel(raw_path)
            
            df = df.where(pd.notnull(df), None)
            
            records = []
            for _, row in df.iterrows():
                rec = {}
                for col, val in row.items():
                    key = clean_excel_key(col)
                    rec[key] = str(val) if val is not None else ""
                records.append(post_process_record(rec))
                
            finalize_document_processing(doc_id, "completed", {"records": records}, 1.0, db)
            return

        # 1b. Handle Word Ingestion (.docx)
        is_word = ext in [".docx", ".doc"]
        if is_word:
            if ext == ".doc":
                raise ValueError("Legacy Word (.doc) format is not supported. Please save/convert the document to modern Word (.docx) format.")
            cache_service.set(f"job:{doc_id}:stage", "ocr", expire_seconds=3600)
            import docx
            try:
                doc_obj = docx.Document(raw_path)
            except Exception as e:
                raise ValueError(f"Failed to parse Word document: {str(e)}. Make sure it is a valid .docx file.")
                
            paragraphs_text = [p.text for p in doc_obj.paragraphs]
            table_text = []
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        table_text.append(cell.text)
            
            full_text = "\n".join(paragraphs_text + table_text)
            ocr_res = {"raw_text": full_text}
            doc.ocr_raw_text = full_text
            db.commit()
            
            cache_service.set(f"job:{doc_id}:stage", "extraction", expire_seconds=3600)
            extracted_data = slm_engine.extract_fields(ocr_res)
            
            conf = 0.0
            if extracted_data:
                non_null_count = sum(1 for v in extracted_data.values() if v is not None)
                conf = non_null_count / len(extracted_data) if len(extracted_data) > 0 else 0.0
                
            finalize_document_processing(doc_id, "completed", extracted_data, conf, db)
            return

        # 1c. Handle PDF Ingestion
        if ext == ".pdf":
            import pypdfium2 as pdfium
            pdf = pdfium.PdfDocument(raw_path)
            if len(pdf) > 0:
                if len(pdf) == 1:
                    # Single-page PDF: process as a flat dictionary
                    # Check OCR Cache (Rule 4)
                    ocr_cache_key = f"ocr:{file_hash}" if file_hash else None
                    cached_ocr_res = cache_service.get(ocr_cache_key) if ocr_cache_key else None
                    
                    cache_service.set(f"job:{doc_id}:stage", "ocr", expire_seconds=3600)
                    if cached_ocr_res:
                        ocr_res = cached_ocr_res
                        append_admin_log("deduplication", "INFO", f"OCR Cache Hit for single-page PDF '{doc.filename}'. Restoring cached OCR data.")
                    else:
                        page = pdf[0]
                        pil_image = page.render(scale=2).to_pil()
                        pdf_image_filename = f"pdf_page_{doc.id}.png"
                        pdf_image_path = os.path.join(settings.UPLOAD_DIR, pdf_image_filename)
                        pil_image.save(pdf_image_path)
                        
                        doc.filename = pdf_image_filename
                        db.commit()
                        
                        ocr_res = ocr_engine.extract_text(pdf_image_path)
                        if ocr_cache_key:
                            cache_service.set(ocr_cache_key, ocr_res, expire_seconds=24 * 3600) # 24 hr TTL
                    
                    doc.ocr_raw_text = ocr_res["raw_text"]
                    db.commit()
                    
                    cache_service.set(f"job:{doc_id}:stage", "extraction", expire_seconds=3600)
                    extracted_data = slm_engine.extract_fields(ocr_res, file_hash=file_hash)
                    
                    non_null_count = sum(1 for v in extracted_data.values() if v is not None)
                    conf = non_null_count / len(extracted_data) if len(extracted_data) > 0 else 0.0
                    finalize_document_processing(doc_id, "completed", extracted_data, conf, db)
                else:
                    # Multi-page PDF: process each page as a separate record
                    records = []
                    all_raw_text = []
                    for idx in range(len(pdf)):
                        cache_service.set(f"job:{doc_id}:stage", "ocr", expire_seconds=3600)
                        # Look up page-level OCR cache
                        page_cache_key = f"ocr:{file_hash}_page_{idx+1}" if file_hash else None
                        cached_page_ocr = cache_service.get(page_cache_key) if page_cache_key else None
                        
                        if cached_page_ocr:
                            ocr_res = cached_page_ocr
                            append_admin_log("deduplication", "INFO", f"OCR Cache Hit for multi-page PDF page {idx+1} of '{doc.filename}'.")
                        else:
                            page = pdf[idx]
                            pil_image = page.render(scale=2).to_pil()
                            pdf_image_filename = f"pdf_page_{doc.id}_page_{idx+1}.png"
                            pdf_image_path = os.path.join(settings.UPLOAD_DIR, pdf_image_filename)
                            pil_image.save(pdf_image_path)
                            
                            if idx == 0:
                                doc.filename = pdf_image_filename
                                db.commit()
                                
                            ocr_res = ocr_engine.extract_text(pdf_image_path)
                            if page_cache_key:
                                cache_service.set(page_cache_key, ocr_res, expire_seconds=24 * 3600)
                                
                        all_raw_text.append(f"--- PAGE {idx+1} ---\n{ocr_res['raw_text']}")
                        
                        cache_service.set(f"job:{doc_id}:stage", "extraction", expire_seconds=3600)
                        extracted_page_data = slm_engine.extract_fields(
                            ocr_res, 
                            file_hash=f"{file_hash}_page_{idx+1}" if file_hash else None
                        )
                        if "records" in extracted_page_data and isinstance(extracted_page_data["records"], list):
                            records.extend(extracted_page_data["records"])
                        else:
                            records.append(extracted_page_data)
                            
                    doc.ocr_raw_text = "\n\n".join(all_raw_text)
                    db.commit()
                    
                    extracted_data = {"records": records}
                    
                    conf = 0.0
                    if records:
                        first_rec = records[0]
                        non_null_count = sum(1 for v in first_rec.values() if v is not None)
                        conf = non_null_count / len(first_rec) if len(first_rec) > 0 else 1.0
                    finalize_document_processing(doc_id, "completed", extracted_data, conf, db)
                return
            else:
                raise ValueError("PDF file is empty")
        else:
            # 3. Check OCR Cache (Rule 4)
            ocr_cache_key = f"ocr:{file_hash}" if file_hash else None
            cached_ocr_res = cache_service.get(ocr_cache_key) if ocr_cache_key else None
            
            cache_service.set(f"job:{doc_id}:stage", "ocr", expire_seconds=3600)
            if cached_ocr_res:
                ocr_res = cached_ocr_res
                append_admin_log("deduplication", "INFO", f"OCR Cache Hit for image '{doc.filename}'. Restoring cached OCR data.")
            else:
                # 2. Run OpenCV Preprocessing (images only)
                preprocessed_filename = f"processed_{os.path.basename(raw_path)}"
                preprocessed_path = os.path.join(settings.UPLOAD_DIR, preprocessed_filename)
                try:
                    ImagePreprocessor.preprocess(raw_path, preprocessed_path)
                    ocr_target_path = preprocessed_path
                except Exception as preprocess_err:
                    print(f"Preprocessing warning: {preprocess_err}. Proceeding with raw file.")
                    ocr_target_path = raw_path
                
                # 3. OCR Extraction
                ocr_res = ocr_engine.extract_text(ocr_target_path)
                if ocr_cache_key:
                    cache_service.set(ocr_cache_key, ocr_res, expire_seconds=24 * 3600) # 24 hr TTL
            
        doc.ocr_raw_text = ocr_res["raw_text"]
        db.commit()
        
        # 4. SLM Structured Extraction
        cache_service.set(f"job:{doc_id}:stage", "extraction", expire_seconds=3600)
        extracted_data = slm_engine.extract_fields(ocr_res, file_hash=file_hash)
        
        # Calculate confidence score
        conf = 0.0
        if "records" in extracted_data and isinstance(extracted_data["records"], list) and len(extracted_data["records"]) > 0:
            first_rec = extracted_data["records"][0]
            non_null_count = sum(1 for v in first_rec.values() if v is not None)
            conf = non_null_count / len(first_rec) if len(first_rec) > 0 else 1.0
        else:
            non_null_count = sum(1 for v in extracted_data.values() if v is not None)
            conf = non_null_count / len(extracted_data) if len(extracted_data) > 0 else 0.0
        
        finalize_document_processing(doc_id, "completed", extracted_data, conf, db)
        
    except Exception as e:
        print(f"Background task failed for document {doc_id}: {e}")
        finalize_document_processing(doc_id, "failed", None, 0.0, db)
    finally:
        db.close()


# Simple global in-memory deduplication cache
# Maps SHA-256 hash -> document ID
DEDUPLICATION_CACHE = {}

@router.post("/upload", response_model=DocumentResponse)
def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    import hashlib
    from .admin import append_admin_log, load_admin_data

    # 1. Load active security settings from cache (Point 10)
    allowed_exts = cache_service.get("config:allowed_extensions")
    if allowed_exts is None:
        admin_settings = load_admin_data()["settings"]
        allowed_exts = admin_settings.get("allowed_extensions", ["pdf", "png", "jpeg", "jpg", "tiff"])
        cache_service.set("config:allowed_extensions", allowed_exts, expire_seconds=600)
    else:
        admin_settings = load_admin_data()["settings"]
        
    max_size_mb = admin_settings.get("max_file_size_mb", 20)
    dup_detection = admin_settings.get("duplicate_detection_sha256", True)

    # 2. File Upload Security: Extension validation
    filename = file.filename
    ext = os.path.splitext(filename.lower())[1].lstrip(".")
    if ext not in allowed_exts:
        append_admin_log("quarantine", "WARNING", f"Rejected upload of '{filename}' (type: .{ext}). Profile not allowed under settings.")
        raise HTTPException(
            status_code=400,
            detail=f"Security Policy Violation: File type .{ext} is restricted. Allowed profiles: {', '.join(allowed_exts)}."
        )

    # 3. Read bytes to compute size and SHA-256 hash
    file_bytes = file.file.read()
    file_size_mb = len(file_bytes) / (1024 * 1024)
    
    # File size validation
    if file_size_mb > max_size_mb:
        append_admin_log("quarantine", "WARNING", f"Rejected upload of '{filename}' (size: {file_size_mb:.2f}MB). Exceeded {max_size_mb}MB limit.")
        raise HTTPException(
            status_code=400,
            detail=f"Security Policy Violation: File exceeds maximum allowed size of {max_size_mb}MB."
        )
        
    # Reset file read pointer
    file.file.seek(0)
    
    # 4. Duplicate Detection via SHA-256 (Rule 6)
    file_hash = hashlib.sha256(file_bytes).hexdigest()
    dup_cache_key = f"dup:hash:{file_hash}"
    cached_doc_id = cache_service.get(dup_cache_key) if dup_detection else None
    if cached_doc_id:
        cached_doc = db.query(DBModelDocument).filter(DBModelDocument.id == cached_doc_id).first()
        if cached_doc and cached_doc.status == "completed":
            append_admin_log("deduplication", "INFO", f"Duplicate SHA-256 match for '{filename}'. Returning cached document ID: {cached_doc_id[:8]}...")
            return cached_doc

    # Save file
    file_path = os.path.join(settings.UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        buffer.write(file_bytes)
        
    # Create DB Record
    db_doc = DBModelDocument(
        filename=file.filename,
        storage_path=file_path,
        mime_type=file.content_type,
        status="pending"
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    
    # Cache the hash for deduplication (30 days TTL)
    if file_hash:
        cache_service.set(dup_cache_key, db_doc.id, expire_seconds=30 * 24 * 3600)
    
    # Initialize Job progress stage in cache
    cache_service.set(f"job:{db_doc.id}:stage", "pending", expire_seconds=3600)

    # 5. Live logging and virus scan scheduling
    append_admin_log("security", "INFO", f"Ingested '{filename}'. Scheduling ClamAV quarantine virus scan...")

    # Trigger background pipeline
    from ..database import SessionLocal
    background_tasks.add_task(process_document_task, db_doc.id, SessionLocal)
    
    return db_doc


# --- Async Job API endpoints (Point 9) ---

@router.post("/jobs/upload")
def upload_document_job(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    """
    Asynchronously uploads file, returns job ID immediately, and runs parsing in the background.
    """
    # Simply reuse upload logic, but return a lighter job response payload
    doc = upload_document(background_tasks, file, db)
    return {
        "job_id": doc.id,
        "status": doc.status,
        "created_at": doc.created_at.isoformat() if doc.created_at else datetime.utcnow().isoformat() + "Z"
    }

@router.get("/jobs/{job_id}/status")
def get_job_status(job_id: str, db: Session = Depends(get_db)):
    """
    Retrieves the status and specific progress stage of an extraction job.
    """
    doc = db.query(DBModelDocument).filter(DBModelDocument.id == job_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Job not found")
    
    progress_stage = cache_service.get(f"job:{job_id}:stage") or doc.status
    return {
        "job_id": job_id,
        "status": doc.status,
        "progress_stage": progress_stage
    }

@router.get("/jobs/{job_id}/result")
def get_job_result(job_id: str, db: Session = Depends(get_db)):
    """
    Retrieves the parsed/extracted JSON output for completed jobs.
    """
    doc = db.query(DBModelDocument).filter(DBModelDocument.id == job_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Job not found")

    if doc.status in ["pending", "processing"]:
        return JSONResponse(
            status_code=202,
            content={
                "job_id": job_id,
                "status": doc.status,
                "message": "Extraction is still in progress. Please check status again."
            }
        )
        
    return {
        "job_id": job_id,
        "status": doc.status,
        "extracted_json": doc.extracted_json,
        "confidence_score": doc.confidence_score
    }


@router.get("", response_model=List[DocumentResponse])
def get_all_documents(db: Session = Depends(get_db)):
    """Returns a list of all documents, ordered by creation date descending."""
    return db.query(DBModelDocument).order_by(DBModelDocument.created_at.desc()).all()


@router.get("/{doc_id}", response_model=DocumentResponse)
def get_document(doc_id: str, db: Session = Depends(get_db)):
    # Point 14: Frequently Downloaded result cache (10 min TTL)
    doc_cache_key = f"doc:result:{doc_id}"
    cached_doc = cache_service.get(doc_cache_key)
    if cached_doc:
        return cached_doc
        
    doc = db.query(DBModelDocument).filter(DBModelDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    doc_data = {
        "id": doc.id,
        "filename": doc.filename,
        "storage_path": doc.storage_path,
        "mime_type": doc.mime_type,
        "status": doc.status,
        "extracted_json": doc.extracted_json,
        "corrected_json": doc.corrected_json,
        "confidence_score": doc.confidence_score,
        "ocr_raw_text": doc.ocr_raw_text,
        "created_at": doc.created_at.isoformat() if doc.created_at else None
    }
    
    if doc.status in ["completed", "failed"]:
        cache_service.set(doc_cache_key, doc_data, expire_seconds=600)
    return doc


@router.delete("/{doc_id}")
def delete_document(doc_id: str, db: Session = Depends(get_db)):
    """Deletes a document record and clean up associated files on disk."""
    doc = db.query(DBModelDocument).filter(DBModelDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Invalidate document cache (Point 14)
    cache_service.delete(f"doc:result:{doc_id}")
    
    # Try deleting physical files
    try:
        if doc.storage_path and os.path.exists(doc.storage_path):
            os.remove(doc.storage_path)
            
        # Try deleting preprocessed path
        preprocessed_filename = f"processed_{os.path.basename(doc.storage_path)}"
        preprocessed_path = os.path.join(settings.UPLOAD_DIR, preprocessed_filename)
        if os.path.exists(preprocessed_path):
            os.remove(preprocessed_path)
            
        # Try deleting pdf rendered page image
        pdf_image_filename = f"pdf_page_{doc.id}.png"
        pdf_image_path = os.path.join(settings.UPLOAD_DIR, pdf_image_filename)
        if os.path.exists(pdf_image_path):
            os.remove(pdf_image_path)
            
        # Try deleting screenshots
        screenshot_filename = f"screenshot_{doc.id}.png"
        screenshot_path = os.path.join(settings.UPLOAD_DIR, "screenshots", screenshot_filename)
        if os.path.exists(screenshot_path):
            os.remove(screenshot_path)
    except Exception as e:
        print(f"Warning: Failed to delete physical files for document {doc_id}: {e}")
        
    db.delete(doc)
    db.commit()
    return {"message": "Document deleted successfully"}

@router.post("/{doc_id}/review", response_model=DocumentResponse)
def review_document(
    doc_id: str,
    review_data: DocumentReview,
    db: Session = Depends(get_db)
):
    doc = db.query(DBModelDocument).filter(DBModelDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    # 1. Update corrected values
    doc.corrected_json = review_data.corrected_json
    db.commit()
    
    # Invalidate cached document result (Point 14)
    cache_service.delete(f"doc:result:{doc_id}")
    
    # 2. Learn field-to-key mapping memory
    # We compare original extracted labels with user corrected labels
    # If the user corrects/adds value in a field, we look at the raw OCR text to find the label
    # Here we learn mapping corrections: source label (OCR label) -> normalized target key
    # For instance, if user matches 'user_mail' input to 'email' field, we learn it.
    # In review, we are validating extracted fields. We can check if any field values correspond to
    # specific raw OCR words, and increment mapping frequency count.
    raw_text = doc.ocr_raw_text or ""
    
    for key, val in review_data.corrected_json.items():
        if val and isinstance(val, str):
            # Look up if this value exists in raw OCR text and trace nearby words
            # To simplify, we also allow the client to send mapping correlations, 
            # but we can auto-associate here.
            # Find lines in raw OCR containing the value
            lines = raw_text.split('\n')
            for line in lines:
                if val.lower() in line.lower():
                    # Parse potential label (words preceding the value, e.g., "Mail ID: test@gmail.com")
                    parts = line.split(val)
                    if parts and parts[0].strip():
                        source_label = parts[0].strip().rstrip(":").rstrip("-").strip()
                        if len(source_label) > 1 and len(source_label) < 30:
                            # Save in mapping memory
                            # Upsert mapping memory
                            mapping = db.query(DBModelMappingMemory).filter(
                                DBModelMappingMemory.source_label == source_label,
                                DBModelMappingMemory.target_key == key
                            ).first()
                            
                            if mapping:
                                mapping.frequency_count += 1
                            else:
                                mapping = DBModelMappingMemory(
                                    source_label=source_label,
                                    target_key=key,
                                    frequency_count=1
                                )
                                db.add(mapping)
                            db.commit()
                            break

    db.refresh(doc)
    return doc
