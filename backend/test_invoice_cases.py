import os
import sys
import docx
from sqlalchemy.orm import Session

# Add absolute backend path for test imports
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from app.database import SessionLocal, engine, Base
from app.models.document import DBModelDocument
Base.metadata.create_all(bind=engine)
from app.routers.documents import process_document_task
from app.config import settings

def create_positive_invoice(path):
    doc = docx.Document()
    doc.add_heading("INVOICE", 0)
    doc.add_paragraph("Vendor Name: Global Services Ltd")
    doc.add_paragraph("Invoice No: INV-2026-905")
    doc.add_paragraph("Date: 05-06-2026")
    doc.add_paragraph("Total Amount: 14,500.00")
    doc.add_paragraph("Mobile Number: 9876543110")
    doc.add_paragraph("Email: biller@global.com")
    doc.add_paragraph("Item Description: Gold Bangle")
    doc.add_paragraph("Gross Weight: 16.540 g")
    doc.add_paragraph("Stone Weight: 1.200 g")
    doc.add_paragraph("Net Weight: 15.340 g")
    doc.add_paragraph("Purity: 22K (916 hallmark)")
    doc.add_paragraph("Rate per Gram: 7100.00")
    doc.add_paragraph("Making Charges: 450.00")
    doc.add_paragraph("Wastage: 7.5 %")
    doc.add_paragraph("Quantity: 5")
    doc.add_paragraph("Unit Price: 250.00")
    doc.add_paragraph("Discount: 50.00")
    doc.add_paragraph("HSN Code: 5208")
    doc.add_paragraph("Shipping Charges: 120.00")
    doc.add_paragraph("SKU Code: SKU-TEXT-452")
    doc.add_paragraph("Patient Name: Rajesh Kumar")
    doc.add_paragraph("Doctor Name: Dr. Aditi Sharma")
    doc.add_paragraph("Admission Date: 10-06-2026")
    doc.add_paragraph("Discharge Date: 15-06-2026")
    doc.add_paragraph("Room Number: Room 402")
    doc.add_paragraph("Medicine Cost: 4,500.00")
    doc.add_paragraph("Insurance Provider: Star Health")
    doc.add_paragraph("PNR No: PNR1234567")
    doc.add_paragraph("Journey Date: 20-06-2026")
    doc.add_paragraph("Source Location: Chennai Central")
    doc.add_paragraph("Destination Location: Bangalore City")
    doc.add_paragraph("Seat Number: Coach A1, Seat 24")
    doc.add_paragraph("Vehicle No: 12623")
    doc.save(path)
    print(f"Created Positive Invoice at: {path}")

def create_negative_empty(path):
    doc = docx.Document()
    # No text added at all
    doc.save(path)
    print(f"Created Negative Empty Invoice at: {path}")

def create_negative_garbage(path):
    doc = docx.Document()
    doc.add_paragraph("asdfhjk1234908")
    doc.add_paragraph("xyz9876543210qwert")
    doc.add_paragraph("lorem ipsum dolor sit amet")
    doc.save(path)
    print(f"Created Negative Garbage Invoice at: {path}")

def create_negative_partial(path):
    doc = docx.Document()
    doc.add_heading("Partial Document", 0)
    # Only full name, missing vendor name, date, invoice number, amount
    doc.add_paragraph("Full Name: Olivia Wilson")
    doc.save(path)
    print(f"Created Negative Partial Invoice at: {path}")

def run_test_for_file(filename, setup_func):
    doc_path = os.path.join(settings.UPLOAD_DIR, filename)
    setup_func(doc_path)
    
    db = SessionLocal()
    try:
        # Insert DB record
        db_doc = DBModelDocument(
            filename=filename,
            storage_path=doc_path,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            status="pending"
        )
        db.add(db_doc)
        db.commit()
        doc_id = db_doc.id
    finally:
        db.close()
        
    print(f"\n--- Processing {filename} (ID: {doc_id}) ---")
    process_document_task(doc_id, SessionLocal)
    
    db = SessionLocal()
    try:
        db_doc = db.query(DBModelDocument).filter(DBModelDocument.id == doc_id).first()
        print(f"Status: {db_doc.status}")
        print(f"Extracted JSON:\n{db_doc.extracted_json}")
        return db_doc.extracted_json
    finally:
        db.close()

def main():
    print("Starting Invoice Extraction Test Suite...")
    
    # Force fallback to rule-based heuristics for deterministic matching verification
    from app.routers.documents import slm_engine
    original_url = slm_engine.base_url
    slm_engine.base_url = "http://localhost:9999" # Redirect to dead port
    
    try:
        # 1. Positive Case
        pos_json = run_test_for_file("pos_invoice.docx", create_positive_invoice)
        # Assertions for positive case
        assert pos_json.get("vendor_name") == "Global Services Ltd", f"Expected 'Global Services Ltd', got {pos_json.get('vendor_name')}"
        assert pos_json.get("invoice_number") == "INV-2026-905", f"Expected 'INV-2026-905', got {pos_json.get('invoice_number')}"
        assert pos_json.get("invoice_date") == "05-06-2026", f"Expected '05-06-2026', got {pos_json.get('invoice_date')}"
        assert pos_json.get("total_amount") == "14,500.00", f"Expected '14,500.00', got {pos_json.get('total_amount')}"
        assert pos_json.get("mobile_number") == "9876543110", f"Expected '9876543110', got {pos_json.get('mobile_number')}"
        assert pos_json.get("email") == "biller@global.com", f"Expected 'biller@global.com', got {pos_json.get('email')}"
        assert pos_json.get("item_description") == "Gold Bangle", f"Expected 'Gold Bangle', got {pos_json.get('item_description')}"
        assert pos_json.get("gross_weight") == "16.540 g", f"Expected '16.540 g', got {pos_json.get('gross_weight')}"
        assert pos_json.get("stone_weight") == "1.200 g", f"Expected '1.200 g', got {pos_json.get('stone_weight')}"
        assert pos_json.get("net_weight") == "15.340 g", f"Expected '15.340 g', got {pos_json.get('net_weight')}"
        assert pos_json.get("purity") == "22K (916 hallmark)", f"Expected '22K (916 hallmark)', got {pos_json.get('purity')}"
        assert pos_json.get("rate_per_gram") == "7100.00", f"Expected '7100.00', got {pos_json.get('rate_per_gram')}"
        assert pos_json.get("making_charges") == "450.00", f"Expected '450.00', got {pos_json.get('making_charges')}"
        assert pos_json.get("wastage") == "7.5 %", f"Expected '7.5 %', got {pos_json.get('wastage')}"
        
        # New assertions for Textile, E-commerce, Medical, and Booking
        assert pos_json.get("quantity") == "5", f"Expected '5', got {pos_json.get('quantity')}"
        assert pos_json.get("unit_price") == "250.00", f"Expected '250.00', got {pos_json.get('unit_price')}"
        assert pos_json.get("discount") == "50.00", f"Expected '50.00', got {pos_json.get('discount')}"
        assert pos_json.get("hsn_code") == "5208", f"Expected '5208', got {pos_json.get('hsn_code')}"
        assert pos_json.get("shipping_charges") == "120.00", f"Expected '120.00', got {pos_json.get('shipping_charges')}"
        assert pos_json.get("sku_code") == "SKU-TEXT-452", f"Expected 'SKU-TEXT-452', got {pos_json.get('sku_code')}"
        
        assert pos_json.get("patient_name") == "Rajesh Kumar", f"Expected 'Rajesh Kumar', got {pos_json.get('patient_name')}"
        assert pos_json.get("doctor_name") == "Dr. Aditi Sharma", f"Expected 'Dr. Aditi Sharma', got {pos_json.get('doctor_name')}"
        assert pos_json.get("admission_date") == "10-06-2026", f"Expected '10-06-2026', got {pos_json.get('admission_date')}"
        assert pos_json.get("discharge_date") == "15-06-2026", f"Expected '15-06-2026', got {pos_json.get('discharge_date')}"
        assert pos_json.get("room_number") == "Room 402", f"Expected 'Room 402', got {pos_json.get('room_number')}"
        assert pos_json.get("medicine_cost") == "4,500.00", f"Expected '4,500.00', got {pos_json.get('medicine_cost')}"
        assert pos_json.get("insurance_provider") == "Star Health", f"Expected 'Star Health', got {pos_json.get('insurance_provider')}"
        
        assert pos_json.get("pnr_no") == "PNR1234567", f"Expected 'PNR1234567', got {pos_json.get('pnr_no')}"
        assert pos_json.get("journey_date") == "20-06-2026", f"Expected '20-06-2026', got {pos_json.get('journey_date')}"
        assert pos_json.get("source_location") == "Chennai Central", f"Expected 'Chennai Central', got {pos_json.get('source_location')}"
        assert pos_json.get("destination_location") == "Bangalore City", f"Expected 'Bangalore City', got {pos_json.get('destination_location')}"
        assert pos_json.get("seat_number") == "Coach A1, Seat 24", f"Expected 'Coach A1, Seat 24', got {pos_json.get('seat_number')}"
        assert pos_json.get("vehicle_no") == "12623", f"Expected '12623', got {pos_json.get('vehicle_no')}"

        print(">>> Positive Invoice assertions PASSED!")

        # 2. Negative Case 1: Empty
        empty_json = run_test_for_file("neg_empty.docx", create_negative_empty)
        # Check that it extracted empty dictionary or handled gracefully
        assert empty_json == {}, f"Expected empty dict, got {empty_json}"
        print(">>> Empty Invoice assertions PASSED!")

        # 3. Negative Case 2: Garbage
        garbage_json = run_test_for_file("neg_garbage.docx", create_negative_garbage)
        # Check that it didn't extract standard invoice fields out of garbage
        assert "invoice_number" not in garbage_json
        assert "total_amount" not in garbage_json
        print(">>> Garbage Invoice assertions PASSED!")

        # 4. Negative Case 3: Partial
        partial_json = run_test_for_file("neg_partial.docx", create_negative_partial)
        assert partial_json.get("full_name") == "Olivia Wilson"
        assert "invoice_number" not in partial_json
        assert "total_amount" not in partial_json
        print(">>> Partial Invoice assertions PASSED!")

        print("\nAll invoice extraction tests passed successfully!")
    finally:
        # Restore URL
        slm_engine.base_url = original_url

if __name__ == "__main__":
    main()
